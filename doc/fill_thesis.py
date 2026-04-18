#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Fill the thesis template with content from the LaTeX draft.
Uses element references instead of paragraph indices to handle insertions correctly."""

import json, os
from docx import Document
from docx.shared import Pt, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

def clear_paragraph(p):
    for child in list(p._element):
        if child.tag.endswith('}r'):
            p._element.remove(child)

def add_run(p, text, bold=None, size=None):
    run = p.add_run(text)
    if bold is not None:
        run.font.bold = bold
    if size is not None:
        run.font.size = size
    return run

def insert_paragraph_after(paragraph):
    new_p = OxmlElement('w:p')
    paragraph._element.addnext(new_p)
    return Paragraph(new_p, paragraph._element.getparent())

def clear_table(table):
    tbl = table._element
    tbl.getparent().remove(tbl)

def build_items(sections):
    items = []
    for sec in sections:
        items.append((sec['title'], True, sec['title_size']))
        for para in sec['paras']:
            items.append((para, False, 12))
    return items


def main():
    content_path = os.path.join(os.path.dirname(__file__), 'thesis_content.json')
    with open(content_path, 'r', encoding='utf-8') as f:
        C = json.load(f)

    doc = Document('doc/\u8bba\u6587\u6a21\u677f.docx')

    # === Capture element references before any modifications ===
    # These are stable references to the XML elements
    p_abstract_body = [doc.paragraphs[i] for i in [28, 29, 30]]
    p_keywords_cn = doc.paragraphs[32]
    p_abstract_en = doc.paragraphs[35]
    p_keywords_en = doc.paragraphs[37]
    p_ch1_heading = doc.paragraphs[38]
    p_ref_heading = doc.paragraphs[107]
    p_ref_entries = [doc.paragraphs[i] for i in range(108, 127)]  # 19 entries
    p_ack_heading = doc.paragraphs[127]
    p_ack_body = [doc.paragraphs[i] for i in [128, 129]]
    p_appA_heading = doc.paragraphs[132]
    p_appA_body = [doc.paragraphs[i] for i in [133, 134]]
    p_appB_heading = doc.paragraphs[136]
    p_appB_body = [doc.paragraphs[i] for i in [137, 138]]

    # All body paragraphs between ch1 heading and ref heading
    body_paras = [doc.paragraphs[i] for i in range(39, 107)]  # 68 paragraphs

    # === Chinese Abstract ===
    for i, p in enumerate(p_abstract_body):
        clear_paragraph(p)
        if i < len(C['abstract_cn']):
            add_run(p, C['abstract_cn'][i], size=Emu(133350))

    clear_paragraph(p_keywords_cn)
    add_run(p_keywords_cn, '\u5173\u952e\u8bcd\uff1a', bold=True, size=Pt(12))
    add_run(p_keywords_cn, C['keywords_cn'])

    # === English Abstract ===
    clear_paragraph(p_abstract_en)
    add_run(p_abstract_en, ' '.join(C['abstract_en']), size=Pt(12))

    clear_paragraph(p_keywords_en)
    add_run(p_keywords_en, 'Key words: ', bold=True, size=Pt(12))
    add_run(p_keywords_en, C['keywords_en'], size=Pt(12))
    p_keywords_en.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # === Build all chapter content ===
    ch1_items = build_items(C['ch1'])
    ch2_full = [('2 \u6587\u732e\u7efc\u8ff0', True, 16)] + build_items(C['ch2'])
    ch3_full = [('3 \u6570\u636e\u4e0e\u95ee\u9898\u5b9a\u4e49', True, 16)] + build_items(C['ch3'])
    ch4_full = [('4 \u65b9\u6cd5', True, 16)] + build_items(C['ch4'])
    ch5_full = [('5 \u5b9e\u9a8c\u4e0e\u7ed3\u679c', True, 16)] + build_items(C['ch5'])
    ch6_full = build_items(C['ch6'])
    ch7_full = build_items(C['ch7'])
    all_body = ch1_items + ch2_full + ch3_full + ch4_full + ch5_full + ch6_full + ch7_full

    # === Fill body paragraphs ===
    # First, fill existing paragraphs
    for idx, p in enumerate(body_paras):
        if idx < len(all_body):
            text, bold, size_pt = all_body[idx]
            clear_paragraph(p)
            add_run(p, text, bold=bold, size=Pt(size_pt))
        else:
            clear_paragraph(p)

    # If we need more paragraphs, insert after the last body paragraph (before ref heading)
    if len(all_body) > len(body_paras):
        last_p = body_paras[-1] if body_paras else p_ch1_heading
        for idx in range(len(body_paras), len(all_body)):
            text, bold, size_pt = all_body[idx]
            new_p = insert_paragraph_after(last_p)
            add_run(new_p, text, bold=bold, size=Pt(size_pt))
            last_p = new_p

    # === References ===
    refs = C['references']
    for idx, p in enumerate(p_ref_entries):
        if idx < len(refs):
            clear_paragraph(p)
            add_run(p, refs[idx])
        else:
            clear_paragraph(p)

    # If we need more reference paragraphs (we have 39 refs but only 19 slots)
    if len(refs) > len(p_ref_entries):
        last_p = p_ref_entries[-1]
        for idx in range(len(p_ref_entries), len(refs)):
            new_p = insert_paragraph_after(last_p)
            add_run(new_p, refs[idx])
            last_p = new_p

    # === Acknowledgements ===
    ack_text = C.get('acknowledgement', '')
    if ack_text:
        for p in p_ack_body:
            clear_paragraph(p)
        add_run(p_ack_body[0], ack_text, size=Pt(12))

    # === Appendix A ===
    for p in p_appA_body:
        clear_paragraph(p)
    add_run(p_appA_body[0], '\uff08\u8bba\u6587\u68c0\u6d4b\u62a5\u544a\u5c06\u5728\u6700\u7ec8\u63d0\u4ea4\u65f6\u9644\u4e0a\uff09', size=Pt(12))

    # === Appendix B ===
    for p in p_appB_body:
        clear_paragraph(p)
    add_run(p_appB_body[0], C.get('appendix_b', ''), size=Pt(12))

    # === Remove template example tables ===
    for table in list(doc.tables):
        first_cell = table.rows[0].cells[0].text.strip()
        if first_cell in ['\u56fe\u50cf\u7c7b\u522b', '\u7f16\u53f7']:
            clear_table(table)

    # === Save ===
    out = 'doc/\u8bba\u6587_filled.docx'
    doc.save(out)
    print(f'Saved to {out}')

if __name__ == '__main__':
    main()
